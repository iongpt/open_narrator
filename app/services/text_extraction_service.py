"""Text extraction service for various file formats."""

import logging
from pathlib import Path
from typing import cast

logger = logging.getLogger(__name__)


class TextExtractionService:
    """
    Service for extracting text from various file formats.

    Supports: TXT, PDF, EPUB, MOBI, DOCX, DOC, RTF, ODT, HTML, MD
    """

    def extract_text(self, file_path: str) -> str:
        """
        Extract text from a file based on its extension.

        Args:
            file_path: Path to the file to extract text from

        Returns:
            Extracted text content

        Raises:
            ValueError: If file format is not supported
            RuntimeError: If text extraction fails
        """
        path = Path(file_path)
        extension = path.suffix.lower()

        logger.info(f"Extracting text from {extension} file: {path.name}")

        try:
            if extension == ".txt":
                return self._extract_from_txt(path)
            elif extension == ".md":
                return self._extract_from_markdown(path)
            elif extension == ".pdf":
                return self._extract_from_pdf(path)
            elif extension == ".epub":
                return self._extract_from_epub(path)
            elif extension == ".mobi":
                return self._extract_from_mobi(path)
            elif extension == ".docx":
                return self._extract_from_docx(path)
            elif extension == ".doc":
                return self._extract_from_doc(path)
            elif extension == ".rtf":
                return self._extract_from_rtf(path)
            elif extension == ".odt":
                return self._extract_from_odt(path)
            elif extension in (".html", ".htm"):
                return self._extract_from_html(path)
            else:
                raise ValueError(
                    f"Unsupported file format: {extension}. "
                    "Supported formats: .txt, .md, .pdf, .epub, .mobi, .docx, .doc, .rtf, .odt, .html"
                )
        except Exception as e:
            logger.error(f"Failed to extract text from {path.name}: {e}")
            raise RuntimeError(f"Text extraction failed: {str(e)}") from e

    def _extract_from_txt(self, file_path: Path) -> str:
        """
        Extract text from a plain text file.

        Args:
            file_path: Path to the TXT file

        Returns:
            File contents as string
        """
        try:
            # Try UTF-8 first
            with open(file_path, encoding="utf-8") as f:
                text = f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1 if UTF-8 fails
            logger.warning(f"UTF-8 decode failed for {file_path.name}, trying latin-1")
            with open(file_path, encoding="latin-1") as f:
                text = f.read()

        if not text.strip():
            raise ValueError("Text file is empty")

        logger.info(f"Extracted {len(text)} characters from TXT file")
        return text

    def _extract_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from a PDF file using PyPDF2.

        Args:
            file_path: Path to the PDF file

        Returns:
            Extracted text from all pages
        """
        try:
            import PyPDF2
        except ImportError as e:
            raise RuntimeError(
                "PyPDF2 is not installed. Install it with: pip install PyPDF2"
            ) from e

        text_parts = []

        with open(file_path, "rb") as f:
            pdf_reader = PyPDF2.PdfReader(f)
            num_pages = len(pdf_reader.pages)

            logger.info(f"Extracting text from {num_pages} pages")

            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()

                # Guard against None (image/scanned pages)
                if page_text and page_text.strip():
                    text_parts.append(page_text)
                else:
                    logger.warning(
                        f"Page {page_num + 1} has no extractable text (likely image/scanned page)"
                    )

        text = "\n\n".join(text_parts)

        if not text.strip():
            raise ValueError(
                "No text could be extracted from PDF. "
                "This may be a scanned/image-only PDF that requires OCR. "
                "Consider using a PDF with selectable text or an OCR tool first."
            )

        logger.info(f"Extracted {len(text)} characters from PDF")
        return text

    def _extract_from_epub(self, file_path: Path) -> str:
        """
        Extract text from an EPUB file using ebooklib.

        Args:
            file_path: Path to the EPUB file

        Returns:
            Extracted text from all chapters
        """
        try:
            import ebooklib
            from bs4 import BeautifulSoup
            from ebooklib import epub
        except ImportError as e:
            raise RuntimeError(
                "Required libraries not installed. "
                "Install with: pip install ebooklib beautifulsoup4"
            ) from e

        book = epub.read_epub(str(file_path))
        text_parts = []

        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                # Parse HTML content
                soup = BeautifulSoup(item.get_content(), "html.parser")
                item_text = cast(str, soup.get_text(separator="\n", strip=True))
                if item_text:
                    text_parts.append(item_text)

        text: str = "\n\n".join(text_parts)

        if not text.strip():
            raise ValueError("No text could be extracted from EPUB")

        logger.info(f"Extracted {len(text)} characters from EPUB")
        return text

    def _extract_from_mobi(self, file_path: Path) -> str:
        """
        Extract text from a MOBI file.

        MOBI is a proprietary format. This method attempts to use
        mobi library or converts to EPUB first.

        Args:
            file_path: Path to the MOBI file

        Returns:
            Extracted text
        """
        try:
            # Try using mobi library
            import mobi
        except ImportError as e:
            raise RuntimeError(
                "MOBI support requires mobi library. "
                "Install with: pip install mobi\n"
                "Note: MOBI format is proprietary. Consider converting to EPUB first."
            ) from e

        # Extract text from MOBI
        tempdir, filepath = mobi.extract(str(file_path))

        # Ensure cleanup even on errors
        try:
            # Read the extracted HTML files
            text_parts = []
            try:
                from bs4 import BeautifulSoup

                for html_file in Path(tempdir).rglob("*.html"):
                    with open(html_file, encoding="utf-8") as f:
                        soup = BeautifulSoup(f.read(), "html.parser")
                        file_text = cast(str, soup.get_text(separator="\n", strip=True))
                        if file_text:
                            text_parts.append(file_text)
            except ImportError as e:
                raise RuntimeError(
                    "BeautifulSoup4 is required for HTML parsing. "
                    "Install with: pip install beautifulsoup4"
                ) from e

            text: str = "\n\n".join(text_parts)

            if not text.strip():
                raise ValueError("No text could be extracted from MOBI")

            logger.info(f"Extracted {len(text)} characters from MOBI")
            return text

        finally:
            # CRITICAL: Always clean up temp directory
            import shutil

            try:
                if tempdir and Path(tempdir).exists():
                    shutil.rmtree(tempdir)
                    logger.debug(f"Cleaned up MOBI temp directory: {tempdir}")
            except Exception as cleanup_error:
                logger.warning(f"Failed to clean up MOBI temp directory {tempdir}: {cleanup_error}")

    def _extract_from_markdown(self, file_path: Path) -> str:
        """
        Extract text from a Markdown file.

        Args:
            file_path: Path to the MD file

        Returns:
            File contents as string
        """
        # Markdown can be read as plain text for TTS purposes
        return self._extract_from_txt(file_path)

    def _extract_from_docx(self, file_path: Path) -> str:
        """
        Extract text from a DOCX file using python-docx.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text from all paragraphs
        """
        try:
            from docx import Document
        except ImportError as e:
            raise RuntimeError(
                "python-docx is not installed. Install it with: pip install python-docx"
            ) from e

        doc = Document(str(file_path))
        text_parts = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                text_parts.append(para_text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)

        text: str = "\n\n".join(text_parts)

        if not text.strip():
            raise ValueError("No text could be extracted from DOCX")

        logger.info(f"Extracted {len(text)} characters from DOCX")
        return text

    def _extract_from_doc(self, file_path: Path) -> str:
        """
        Extract text from a DOC file (old Word format).

        This is more challenging as DOC is a binary format.
        Uses textract if available, otherwise provides helpful error.

        Args:
            file_path: Path to the DOC file

        Returns:
            Extracted text
        """
        try:
            import textract
        except ImportError as e:
            raise RuntimeError(
                "DOC format requires textract library and system dependencies.\n"
                "Install with: pip install textract\n"
                "System dependencies: antiword (brew install antiword on macOS)\n"
                "Alternative: Convert DOC to DOCX first for better support."
            ) from e

        try:
            text: str = textract.process(str(file_path)).decode("utf-8")

            if not text.strip():
                raise ValueError("No text could be extracted from DOC")

            logger.info(f"Extracted {len(text)} characters from DOC")
            return text

        except Exception as e:
            raise RuntimeError(
                f"Failed to extract from DOC file: {str(e)}\n"
                "Consider converting to DOCX format for better compatibility."
            ) from e

    def _extract_from_rtf(self, file_path: Path) -> str:
        """
        Extract text from an RTF file using striprtf.

        Args:
            file_path: Path to the RTF file

        Returns:
            Extracted text
        """
        try:
            from striprtf.striprtf import rtf_to_text
        except ImportError as e:
            raise RuntimeError(
                "striprtf is not installed. Install it with: pip install striprtf"
            ) from e

        with open(file_path, encoding="utf-8", errors="ignore") as f:
            rtf_content = f.read()

        text: str = cast(str, rtf_to_text(rtf_content))

        if not text.strip():
            raise ValueError("No text could be extracted from RTF")

        logger.info(f"Extracted {len(text)} characters from RTF")
        return text

    def _extract_from_odt(self, file_path: Path) -> str:
        """
        Extract text from an ODT file (OpenDocument Text).

        Args:
            file_path: Path to the ODT file

        Returns:
            Extracted text
        """
        try:
            from odf import text as odf_text
            from odf.opendocument import load
        except ImportError as e:
            raise RuntimeError("odfpy is not installed. Install it with: pip install odfpy") from e

        doc = load(str(file_path))
        text_parts = []

        # Extract all text elements
        import re

        for element in doc.getElementsByType(odf_text.P):
            element_text = str(element)
            # Remove any remaining tags
            element_text = re.sub(r"<[^>]+>", "", element_text)
            if element_text.strip():
                text_parts.append(element_text.strip())

        text: str = "\n\n".join(text_parts)

        if not text.strip():
            raise ValueError("No text could be extracted from ODT")

        logger.info(f"Extracted {len(text)} characters from ODT")
        return text

    def _extract_from_html(self, file_path: Path) -> str:
        """
        Extract text from an HTML file.

        Args:
            file_path: Path to the HTML file

        Returns:
            Extracted text
        """
        try:
            from bs4 import BeautifulSoup
        except ImportError as e:
            raise RuntimeError(
                "BeautifulSoup4 is required for HTML parsing. "
                "Install with: pip install beautifulsoup4"
            ) from e

        with open(file_path, encoding="utf-8", errors="ignore") as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        text: str = cast(str, soup.get_text(separator="\n", strip=True))

        if not text.strip():
            raise ValueError("No text could be extracted from HTML")

        logger.info(f"Extracted {len(text)} characters from HTML")
        return text


def get_text_extraction_service() -> TextExtractionService:
    """
    Get text extraction service instance.

    Returns:
        TextExtractionService instance
    """
    return TextExtractionService()
